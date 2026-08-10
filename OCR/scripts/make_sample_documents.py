"""Build sample DICOM and Word documents carrying known PHI."""
import argparse
import os
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

PATIENT_NAME = "Jane Doe"
PATIENT_ID_VALUE = "MRN4471203"
PHYSICIAN = "Dr Alan Grant"
INSTITUTION = "St Elsewhere General"
BIRTH_DATE = "19780414"
PHONE = "555-0142"


def burn_text(image, text: str, origin=(10, 10)):
    """Draw text into a numpy image, the way a modality burns a banner in."""
    from PIL import Image, ImageDraw

    pil = Image.fromarray(image)
    draw = ImageDraw.Draw(pil)
    draw.text(origin, text, fill=255)
    return __import__("numpy").asarray(pil)


def make_dicom(output_path: Path) -> Path:
    import numpy as np
    import pydicom
    from pydicom.dataset import Dataset, FileMetaDataset
    from pydicom.uid import ExplicitVRLittleEndian, generate_uid

    rows, cols = 256, 512
    pixels = np.full((rows, cols), 40, dtype=np.uint8)

    # Two banners, because burned-in PHI is rarely in one place.
    pixels = burn_text(pixels, f"{PATIENT_NAME}  {PATIENT_ID_VALUE}", (8, 8))
    pixels = burn_text(pixels, f"DOB {BIRTH_DATE}  {INSTITUTION}", (8, 30))

    dataset = Dataset()
    dataset.PatientName = PATIENT_NAME
    dataset.PatientID = PATIENT_ID_VALUE
    dataset.PatientBirthDate = BIRTH_DATE
    dataset.PatientSex = "F"
    dataset.PatientTelephoneNumbers = PHONE
    dataset.ReferringPhysicianName = PHYSICIAN
    dataset.PerformingPhysicianName = PHYSICIAN
    dataset.InstitutionName = INSTITUTION
    dataset.InstitutionAddress = "1 Hospital Way"
    dataset.AccessionNumber = "ACC00099"
    dataset.StudyID = "STUDY-1"
    dataset.StudyDate = "20260101"
    dataset.StudyTime = "101500"
    dataset.Modality = "OT"
    dataset.SeriesInstanceUID = generate_uid()
    dataset.StudyInstanceUID = generate_uid()
    dataset.SOPInstanceUID = generate_uid()
    dataset.SOPClassUID = "1.2.840.10008.5.1.4.1.1.7"

    block = dataset.private_block(0x000B, "HIVE TEST", create=True)
    block.add_new(0x01, "LO", PATIENT_NAME)

    dataset.SamplesPerPixel = 1
    dataset.PhotometricInterpretation = "MONOCHROME2"
    dataset.Rows = rows
    dataset.Columns = cols
    dataset.BitsAllocated = 8
    dataset.BitsStored = 8
    dataset.HighBit = 7
    dataset.PixelRepresentation = 0
    dataset.PixelData = pixels.tobytes()

    meta = FileMetaDataset()
    meta.MediaStorageSOPClassUID = dataset.SOPClassUID
    meta.MediaStorageSOPInstanceUID = dataset.SOPInstanceUID
    meta.TransferSyntaxUID = ExplicitVRLittleEndian
    meta.ImplementationClassUID = generate_uid()
    dataset.file_meta = meta

    output_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        dataset.save_as(str(output_path), enforce_file_format=True)
    except TypeError:
        dataset.save_as(str(output_path), write_like_original=False)
    return output_path


def make_docx(output_path: Path) -> Path:
    import docx

    document = docx.Document()

    document.add_heading("Referral letter", level=1)
    document.add_paragraph(
        f"Patient {PATIENT_NAME}, born {BIRTH_DATE}, reachable on {PHONE}."
    )
    document.add_paragraph(f"Referred by {PHYSICIAN} at {INSTITUTION}.")

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = PATIENT_NAME
    table.cell(1, 0).text = "MRN"
    table.cell(1, 1).text = PATIENT_ID_VALUE

    # Header and footer, where a patient banner usually lives.
    section = document.sections[0]
    section.header.paragraphs[0].text = f"{PATIENT_NAME} - {PATIENT_ID_VALUE}"
    section.footer.paragraphs[0].text = f"Printed at {INSTITUTION}"

    properties = document.core_properties
    properties.author = PHYSICIAN
    properties.last_modified_by = PHYSICIAN
    properties.title = f"Referral for {PATIENT_NAME}"
    properties.comments = f"Contact {PHONE}"

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path


def main(argv=None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output-dir", default="samples")
    parser.add_argument("--kind", choices=("dicom", "docx", "both"), default="both")
    args = parser.parse_args(argv)

    out = Path(args.output_dir).expanduser()

    if args.kind in ("dicom", "both"):
        print(make_dicom(out / "sample_study.dcm"))
    if args.kind in ("docx", "both"):
        print(make_docx(out / "sample_referral.docx"))

    return 0


if __name__ == "__main__":
    sys.exit(main())
