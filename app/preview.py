"""Turn a document into something a browser will actually display.

A PDF an `<iframe>` renders on its own. A DICOM and a Word document it
downloads instead -- which is what made "Original" a download button for
half the files in an application.

Neither format has a browser-native representation, so one is made here:
a DICOM frame becomes a PNG, a Word document becomes its text. Both
happen server-side because the alternative is shipping a DICOM codec and
a docx parser into the bundle, and because the transfer syntaxes a PACS
actually emits need the same decoders `OCR/` needs -- see the note in
OCR/requirements-ocr.txt.

This renders the file as it is. Previewing an original shows the original,
burned-in identifiers and all; that is the same content the download
already served, behind the same permission.
"""
from pathlib import Path
from typing import Dict, List, Tuple

from app.errors import ValidationError
from app.logging_setup import get_logger

log = get_logger(__name__)

# Downscale anything larger. A mammography frame is 4000+ pixels square
# and nobody is inspecting that in a modal.
MAX_DIMENSION = 2400

# Word documents can be long; the viewer is a preview, not an editor.
MAX_PARAGRAPHS = 2000
MAX_TABLE_ROWS = 500


# ------------------------------------------------------------------ dicom


def _to_display_array(dataset, frame_index: int):
    """One frame as 8-bit RGB or greyscale, windowed the way it is meant to be seen."""
    import numpy as np

    try:
        pixels = dataset.pixel_array
    except Exception as exc:
        raise ValidationError(
            f"This DICOM's pixel data could not be decoded: {exc}"
        ) from exc

    frame_count = frame_number(dataset, pixels)

    if frame_index < 0 or frame_index >= frame_count:
        raise ValidationError(
            f"Frame {frame_index} does not exist; this study has {frame_count}"
        )

    samples = int(getattr(dataset, "SamplesPerPixel", 1) or 1)
    if frame_count > 1:
        frame = pixels[frame_index]
    elif pixels.ndim == 3 and samples == 1:
        frame = pixels[frame_index]
    else:
        frame = pixels

    # Window/level, when the study says how it wants to be displayed.
    if samples == 1:
        try:
            from pydicom.pixels import apply_voi_lut

            frame = apply_voi_lut(frame, dataset, index=0)
        except Exception:
            # No LUT, or one this file cannot honour -- fall through to
            # the plain min/max stretch below, which always works.
            pass

    frame = np.asarray(frame)

    if frame.dtype != np.uint8:
        low, high = float(frame.min()), float(frame.max())
        if high > low:
            frame = (frame.astype(np.float64) - low) / (high - low)
        else:
            frame = np.zeros_like(frame, dtype=np.float64)
        frame = (frame * 255.0).astype(np.uint8)

    # MONOCHROME1 counts white as zero; shown as-is it is a negative.
    if str(getattr(dataset, "PhotometricInterpretation", "")) == "MONOCHROME1":
        frame = 255 - frame

    return frame


def frame_number(dataset, pixels=None) -> int:
    """How many frames this study holds."""
    declared = int(getattr(dataset, "NumberOfFrames", 1) or 1)
    if declared > 1:
        return declared

    if pixels is not None:
        samples = int(getattr(dataset, "SamplesPerPixel", 1) or 1)
        if pixels.ndim == 3 and samples == 1:
            return int(pixels.shape[0])

    return 1


def render_dicom_png(path: Path, frame_index: int = 0) -> Tuple[bytes, int]:
    """One frame as PNG bytes, plus the number of frames available."""
    import pydicom
    from PIL import Image

    try:
        dataset = pydicom.dcmread(str(path), force=True)
    except Exception as exc:
        raise ValidationError(f"This file is not readable as DICOM: {exc}") from exc

    if not hasattr(dataset, "PixelData"):
        raise ValidationError("This DICOM carries no image data to display")

    array = _to_display_array(dataset, frame_index)

    image = Image.fromarray(array)
    if image.mode not in ("L", "RGB"):
        image = image.convert("RGB")

    if max(image.size) > MAX_DIMENSION:
        image.thumbnail((MAX_DIMENSION, MAX_DIMENSION))

    from io import BytesIO

    buffer = BytesIO()
    image.save(buffer, format="PNG", optimize=True)

    total = frame_number(dataset, dataset.pixel_array)
    log.info(
        "dicom_rendered",
        path=str(path),
        frame=frame_index,
        frames=total,
        size=image.size,
    )
    return buffer.getvalue(), total


# ------------------------------------------------------------------- word


def _table_rows(table) -> List[List[str]]:
    rows = []
    for row in table.rows[:MAX_TABLE_ROWS]:
        rows.append([cell.text.strip() for cell in row.cells])
    return rows


def read_word_document(path: Path) -> Dict:
    """A Word document's text, as structure rather than markup.

    Deliberately not HTML: this is somebody's uploaded document, and
    handing the browser markup out of it invites the obvious injection.
    The client renders these strings as text nodes.
    """
    import docx

    try:
        document = docx.Document(str(path))
    except Exception as exc:
        raise ValidationError(
            "This Word document could not be opened. Only the 2007+ .docx "
            f"format can be previewed: {exc}"
        ) from exc

    blocks = []
    for paragraph in document.paragraphs[:MAX_PARAGRAPHS]:
        text = paragraph.text.strip()
        if not text:
            continue
        blocks.append(
            {
                "kind": "heading" if paragraph.style.name.startswith("Heading") else "paragraph",
                "style": paragraph.style.name,
                "text": text,
            }
        )

    tables = [_table_rows(table) for table in document.tables]

    log.info(
        "word_document_read",
        path=str(path),
        blocks=len(blocks),
        tables=len(tables),
    )
    return {
        "blocks": blocks,
        "tables": tables,
        "truncated": len(document.paragraphs) > MAX_PARAGRAPHS,
    }
