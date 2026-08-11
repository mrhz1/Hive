"""Extract metadata from an uploaded document."""
import datetime as _datetime
from pathlib import Path
from typing import Any, Dict, Optional, Tuple

from app.filetype import OLE_MAGIC
from app.logging_setup import get_logger
from app.schemas import METADATA_EXTENSIONS

log = get_logger(__name__)

MAX_VALUE_CHARS = 2000
MAX_FIELDS = 200


def file_type_for(extension: str) -> Optional[str]:
    """'pdf' / 'dicom' / 'word', or None when we do not read this format."""
    return METADATA_EXTENSIONS.get((extension or "").lower().lstrip("."))


def _clean(value: Any) -> Optional[str]:
    """Normalise one attribute to a short string, or None to drop it."""
    if value is None:
        return None
    if isinstance(value, (_datetime.datetime, _datetime.date)):
        return value.isoformat()
    if isinstance(value, (bytes, bytearray)):
        # Binary attributes are rarely meaningful and never displayable.
        return None

    text = str(value).strip()
    if not text:
        return None
    return text[:MAX_VALUE_CHARS]


def _collect(pairs) -> Dict[str, str]:
    out: Dict[str, str] = {}
    for name, value in pairs:
        if len(out) >= MAX_FIELDS:
            break
        cleaned = _clean(value)
        if cleaned is not None:
            out[name] = cleaned
    return out


def _extract_pdf(path: Path) -> Dict[str, str]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    info = reader.metadata or {}

    fields = [("page_count", len(reader.pages)), ("encrypted", reader.is_encrypted)]
    fields.extend((str(key).lstrip("/").lower(), value) for key, value in info.items())
    return _collect(fields)


def _extract_dicom(path: Path) -> Dict[str, str]:
    import pydicom

    dataset = pydicom.dcmread(str(path), stop_before_pixels=True, force=True)

    def pairs():
        for element in dataset:
            if element.tag == 0x7FE00010 or element.VR == "SQ":
                continue
            keyword = element.keyword or str(element.tag)
            yield keyword, element.value

    return _collect(pairs())


def _extract_docx(path: Path) -> Dict[str, str]:
    import docx

    document = docx.Document(str(path))
    core = document.core_properties

    return _collect(
        [
            ("title", core.title),
            ("author", core.author),
            ("subject", core.subject),
            ("keywords", core.keywords),
            ("category", core.category),
            ("comments", core.comments),
            ("last_modified_by", core.last_modified_by),
            ("revision", core.revision),
            ("created", core.created),
            ("modified", core.modified),
            ("paragraph_count", len(document.paragraphs)),
            ("table_count", len(document.tables)),
        ]
    )


def _decoded(value: Any) -> Any:
    """olefile hands back bytes for text properties on older documents."""
    if isinstance(value, (bytes, bytearray)):
        return value.decode("utf-8", "replace").strip("\x00") or None
    return value


def _extract_legacy_doc(path: Path) -> Dict[str, str]:
    """A pre-2007 .doc: OLE2, so python-docx cannot open it at all.

    The properties live in the SummaryInformation streams, which is what
    olefile reads. Without it the row records 'failed' with the reason,
    which is what this format did before and is not made worse by it.
    """
    import olefile

    ole = olefile.OleFileIO(str(path))
    try:
        properties = ole.get_metadata()
        names = list(properties.SUMMARY_ATTRIBS) + list(properties.DOCSUM_ATTRIBS)
        return _collect(
            (name, _decoded(getattr(properties, name, None))) for name in names
        )
    finally:
        ole.close()


def _extract_word(path: Path) -> Dict[str, str]:
    """Word, either generation. The container says which, not the name."""
    with open(path, "rb") as handle:
        head = handle.read(len(OLE_MAGIC))

    if head == OLE_MAGIC:
        return _extract_legacy_doc(path)
    return _extract_docx(path)


_EXTRACTORS = {
    "pdf": _extract_pdf,
    "dicom": _extract_dicom,
    "word": _extract_word,
}


def extract(path: Path, extension: str) -> Tuple[str, Dict[str, str], str, Optional[str]]:
    """Read one file's metadata."""
    file_type = file_type_for(extension)
    if file_type is None:
        return (extension or "unknown").lower(), {}, "unsupported", None

    try:
        metadata = _EXTRACTORS[file_type](path)
    except ImportError as exc:
        log.error("file_metadata_parser_missing", file_type=file_type, error=str(exc))
        return file_type, {}, "failed", f"Parser not installed: {exc}"
    except Exception as exc:
        log.warning(
            "file_metadata_extraction_failed",
            file_type=file_type,
            error=str(exc)[:200],
        )
        return file_type, {}, "failed", str(exc)[:500]

    log.info("file_metadata_extracted", file_type=file_type, fields=len(metadata))
    return file_type, metadata, "ok", None
